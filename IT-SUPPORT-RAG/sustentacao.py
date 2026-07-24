import shutil, os
from typing import List, Dict, Any
from dotenv import load_dotenv
import streamlit as st

from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

PERSIST_DIR = "./chroma_clients"
os.makedirs(PERSIST_DIR, exist_ok=True)
os.makedirs("./data", exist_ok=True)  # Temp upload directory

# -------------------------------
# Document Processor
# -------------------------------
class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        )

    def load_documents(self, data_directory: str) -> List[Document]:
        if not os.path.exists(data_directory):
            return []
        loader = DirectoryLoader(data_directory, glob="**/*.*", show_progress=True)
        documents = loader.load()
        return documents

    def split_documents(self, documents: List[Document]) -> List[Document]:
        return self.text_splitter.split_documents(documents)

# -------------------------------
# ChromaDB Manager
# -------------------------------
class ChromaDBManager:
    def __init__(self, persist_directory: str, client_name: str):
        self.persist_directory = persist_directory
        self.client_name = client_name
        self.embedding_function = OpenAIEmbeddings()
        os.makedirs(persist_directory, exist_ok=True)
        self.client = None

    def list_clients(self) -> List[str]:
        try:
            chroma = Chroma(embedding_function=self.embedding_function, persist_directory=self.persist_directory)
            return [c.name for c in chroma._client.list_collections()]
        except Exception as e:
            st.warning(f"Error listing clients: {e}")
            return []


    def get_or_create_collection(self):
        if self.client is None:
            self.client = Chroma(
                collection_name=self.client_name,
                embedding_function=self.embedding_function,
                persist_directory=self.persist_directory,
            )
        return self.client
    
    # ------------------------
    # Rename Client (safe)
    # ------------------------
    def rename_client(self, old_name: str, new_name: str):
        """
        Safely rename a client collection in ChromaDB (LangChain version) and locally.
        Handles cases where the collection doesn't yet exist.
        """
        persist_dir = self.persist_directory

        # Initialize a Chroma connection to access the _client
        try:
            chroma = Chroma(
                embedding_function=self.embedding_function,
                persist_directory=persist_dir
            )
            client = chroma._client
        except Exception as e:
            st.error(f"Error initializing Chroma client: {e}")
            return

        # 1️⃣ List all existing collections
        try:
            collections = [c.name for c in client.list_collections()]
        except Exception as e:
            st.error(f"Error listing collections: {e}")
            return

        # 2️⃣ If collection exists, rename in Chroma
        if old_name in collections:
            try:
                old_collection = client.get_collection(old_name)
                items = old_collection.get()

                # Create new collection
                new_vectorstore = Chroma(
                    collection_name=new_name,
                    persist_directory=persist_dir,
                    embedding_function=self.embedding_function,
                    client=client,
                )

                # Move data only if collection has content
                if items and len(items.get("ids", [])) > 0:
                    new_vectorstore._collection.add(
                        ids=items["ids"],
                        embeddings=items["embeddings"],
                        documents=items["documents"],
                        metadatas=items["metadatas"],
                    )

                # Delete old collection
                client.delete_collection(old_name)

                st.success(f"✅ Renamed Chroma collection '{old_name}' → '{new_name}'")

            except Exception as e:
                st.error(f"Error renaming Chroma collection '{old_name}': {e}")
        else:
            st.info(f"No Chroma collection found for '{old_name}', renaming locally only.")

        # 3️⃣ Rename local folder (if it exists)
        old_path = os.path.join("data", old_name)
        new_path = os.path.join("data", new_name)

        try:
            if os.path.exists(old_path):
                shutil.move(old_path, new_path)
                st.success(f"✅ Renamed local folder '{old_name}' → '{new_name}'")
            else:
                st.info(f"No local folder for '{old_name}', skipping folder rename.")
        except Exception as e:
            st.error(f"Error renaming local folder: {e}")

    def add_documents(self, file_paths: List[str]):
        if not self.client:
            self.get_or_create_collection()
        processor = DocumentProcessor()
        documents = []
        for path in file_paths:
            ext = os.path.splitext(path)[1].lower()
            if ext == ".pdf":
                loader = PyPDFLoader(path)
                documents.extend(loader.load())
            elif ext in [".txt", ".md"]:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                    documents.append(Document(page_content=content, metadata={"source": path}))
            elif ext in [".docx"]:
                try:
                    import docx
                    doc = docx.Document(path)
                    content = "\n".join([p.text for p in doc.paragraphs])
                    documents.append(Document(page_content=content, metadata={"source": path}))
                except Exception as e:
                    st.error(f"Error reading DOCX file {path}: {e}")
            else:
                st.warning(f"Unsupported file type: {path}")

        if documents:
            splits = processor.split_documents(documents)
            self.client.add_documents(splits)
            st.success(f"Added {len(splits)} document chunks to client '{self.client_name}'.")

    def list_documents(self, client_name: str):
        try:
            chroma = Chroma(
                collection_name=client_name,
                embedding_function=self.embedding_function,
                persist_directory=self.persist_directory
            )
            items = chroma.get()
            documents = []

            # Group chunks by their source file
            grouped = {}
            for doc_id, content, metadata in zip(items["ids"], items["documents"], items["metadatas"]):
                source = metadata.get("source", "Unknown")
                if source not in grouped:
                    grouped[source] = []
                grouped[source].append(content)

            # Build a clean document list (combine chunks)
            for source, chunks in grouped.items():
                combined_text = " ".join(chunks)
                documents.append({
                    "source": source,
                    "content": combined_text,
                })

            return documents

        except Exception as e:
            st.warning(f"Error listing documents for {client_name}: {e}")
            return []
    
    def delete_document(self, client_name: str, source_path: str):
        try:
            # Open the collection for that client
            chroma = Chroma(
                collection_name=client_name,
                embedding_function=self.embedding_function,
                persist_directory=self.persist_directory,
            )

            # Get all items
            items = chroma.get()

            # Find IDs of chunks belonging to this source
            ids_to_delete = [
                doc_id for doc_id, meta in zip(items["ids"], items["metadatas"])
                if meta.get("source") == source_path
            ]

            if not ids_to_delete:
                return f"No documents found for source: {source_path}"

            # Delete them from the collection
            chroma.delete(ids=ids_to_delete)

            # Also delete the file from the local data folder
            if os.path.exists(source_path):
                os.remove(source_path)
                return f"Deleted document '{os.path.basename(source_path)}' (and its chunks)"
            else:
                return f"Deleted {len(ids_to_delete)} chunks, but file not found locally"

        except Exception as e:
            return f"Error deleting document: {e}"



    def delete_client(self):
        try:
            chroma = Chroma(embedding_function=self.embedding_function, persist_directory=self.persist_directory)
            chroma._client.delete_collection(self.client_name)
            st.success(f"Deleted client '{self.client_name}'.")
            # Remove from session cache
            if "vector_stores" in st.session_state and self.client_name in st.session_state.vector_stores:
                del st.session_state.vector_stores[self.client_name]

            # Delete associated documents
            client_data_path = os.path.join("data", self.client_name)
            if os.path.exists(client_data_path):
                shutil.rmtree(client_data_path)
                st.info(f"Deleted documents for client '{self.client_name}'.")
        except Exception as e:
            st.error(f"Error deleting client: {e}")

# -------------------------------
# RAG System
# -------------------------------
class RAGSystem:
    def __init__(self, vector_store: Chroma):
        self.vector_store = vector_store
        self.llm = ChatOpenAI(temperature=0, model="gpt-4o-mini")

    def query(self, query: str, top_k: int = 5) -> Dict[str, Any]:
        results = self.vector_store.similarity_search_with_score(query, k=top_k)
        if not results:
            return {"answer": "No relevant documents found.", "sources": []}

        context_parts = []
        sources = []
        for i, (doc, score) in enumerate(results, 1):
            relevance_score = round(score, 3)
            context_parts.append(f"[Document {i}]: {doc.page_content}")
            sources.append({"document": doc, "score": relevance_score})

        context = "\n\n".join(context_parts)

        prompt = PromptTemplate(
            template="""Based on the provided context, answer the question comprehensively.
                        Include relevant quotes and cite documents using [Document X]. If not in context, say so.

                        Context:
                        {context}

                        Question: {question}

                        Answer:""",
            input_variables=["context", "question"]
        )

        response = self.llm.invoke(prompt.format(context=context, question=query))
        return {"answer": response.content, "sources": sources}

# -------------------------------
# Streamlit UI
# -------------------------------


def main():
    st.set_page_config(page_title="RAG Client System", layout="wide")
    st.title("📚Knowledge Base: Client RAG System with GPT-4o-mini")

    # ------------------------
    # Initialize session state
    # ------------------------
    defaults = {
        "vector_stores": {},
        "clients": [],
        "new_client_name": "",
        "active_client": "-- None --",
        "query_input": "",
        "clear_chat_trigger": False,  # Add a flag for safe clearing
        "clear_new_client_trigger" : False,
        "edit_mode": False, # For toggling edit mode visibility
        "delete_mode": False, # Temporary input storage during rename
        "edit_temp_name": ""  #For showing delete confirmation popup

    }
    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default

    # ------------------------
    # BEFORE creating widgets: handle trigger flags (safe clearing)
    # ------------------------
    if st.session_state.clear_chat_trigger:
        st.session_state.query_input = ""
        st.session_state.clear_chat_trigger = False

    if st.session_state.clear_new_client_trigger:
        st.session_state.new_client_name = ""
        st.session_state.clear_new_client_trigger = False

    # ------------------------
    # Helper functions (set triggers instead of clearing directly)
    # ------------------------
    def set_clear_new_client_trigger():
        st.session_state.clear_new_client_trigger = True

    def set_clear_chat_trigger():
        st.session_state.clear_chat_trigger = True


    # ------------------------
    # Sidebar: Client Management
    # ------------------------
    st.sidebar.header("Client Management")

    db_manager = ChromaDBManager(PERSIST_DIR, client_name=None)

    # Load clients if empty
    if not st.session_state.clients:
        st.session_state.clients = db_manager.list_clients()

    selected_client = st.sidebar.selectbox(
        "Select Existing Client",
        ["-- None --"] + st.session_state.clients,
        index=0 if st.session_state.active_client not in st.session_state.clients else
              st.session_state.clients.index(st.session_state.active_client) + 1,
        key="selectbox_client",
        on_change=set_clear_new_client_trigger
    )

    new_client_name = st.sidebar.text_input(
        "Or add new client",
        value=st.session_state.new_client_name,
        key="new_client_name"
    )

    # ------------------------
    # Add New Client
    # ------------------------

    if st.sidebar.button("Add New Client"):
        client_name = new_client_name.strip()
        if client_name and client_name not in st.session_state.clients:
            db_manager.client_name = client_name
            vector_store = db_manager.get_or_create_collection()
            st.session_state.vector_stores[client_name] = vector_store

            st.session_state.clients.append(client_name)
            st.session_state.active_client = client_name

            # Trigger clearing on next rerun (safe)
            st.session_state.clear_new_client_trigger = True

            os.makedirs(os.path.join("data", client_name), exist_ok=True)
            st.rerun()


    # ------------------------
    # Edit / Delete Client Controls
    # ------------------------
    if selected_client != "-- None --":
        st.sidebar.subheader("Manage Client")

        # Two buttons side by side: Edit and Delete
        col1, col2 = st.sidebar.columns(2)

        # --- EDIT BUTTON ---
        with col1:
            if st.button("✏️ Edit", key="edit_client_btn"):
                st.session_state.edit_mode = True
                st.session_state.edit_temp_name = st.session_state.active_client
                st.session_state.delete_mode = False  # ensure not in delete mode
                st.rerun()

        # --- DELETE BUTTON ---
        with col2:
            if st.button("🗑️ Delete", key="delete_client_btn"):
                st.session_state.delete_mode = True
                st.session_state.edit_mode = False
                st.rerun()

        # --- DELETE CONFIRMATION POPUP ---
        if st.session_state.get("delete_mode", False):
            st.sidebar.markdown(
                f"<div style='background-color:#ffe6e6;padding:10px;border-radius:10px'>"
                f"<strong>⚠️ Confirm Deletion:</strong><br>"
                f"Are you sure you want to permanently delete <b>{st.session_state.active_client}</b>?"
                f"</div>",
                unsafe_allow_html=True
            )

            col3, col4 = st.sidebar.columns(2)
            with col3:
                if st.button("✅ Yes, Delete", key="confirm_delete_btn"):
                    client_to_delete = st.session_state.active_client
                    db_manager.client_name = client_to_delete
                    db_manager.delete_client()

                    st.session_state.vector_stores.pop(client_to_delete, None)
                    st.session_state.clients = db_manager.list_clients()
                    st.session_state.active_client = "-- None --"
                    st.session_state.delete_mode = False

                    st.sidebar.success(f"Client '{client_to_delete}' deleted.")
                    st.rerun()

            with col4:
                if st.button("❌ Cancel", key="cancel_delete_btn"):
                    st.session_state.delete_mode = False
                    st.rerun()

        # --- EDIT MODE SECTION ---
        if st.session_state.get("edit_mode", False):
            new_name = st.sidebar.text_input(
                "Enter new name:",
                value=st.session_state.edit_temp_name,
                key="edit_temp_name_input"
            )

            col5, col6 = st.sidebar.columns(2)
            with col5:
                if st.button("💾 Save", key="save_edit_btn"):
                    old_name = st.session_state.active_client
                    updated_name = new_name.strip()

                    if updated_name and updated_name != old_name:
                        # Rename in ChromaDB + filesystem
                        db_manager.rename_client(old_name, updated_name)

                        # Update session state
                        st.session_state.clients = [
                            updated_name if c == old_name else c
                            for c in st.session_state.clients
                        ]
                        st.session_state.active_client = updated_name
                        st.session_state.edit_mode = False
                        st.rerun()

            with col6:
                if st.button("❌ Cancel", key="cancel_edit_btn"):
                    st.session_state.edit_mode = False
                    st.rerun()
    # else:
    #     selected_client = "-- None --"
    #     set_clear_new_client_trigger()
    #     st.session_state.active_client = "-- None --"
    #     st.rerun()
        

    # ------------------------
    # Determine active client
    # ------------------------
    if selected_client != "-- None --":
        st.session_state.active_client = selected_client

    client_name = st.session_state.active_client
    show_interface = client_name != "-- None --"

    # ------------------------
    # Client Interface
    # ------------------------
    if show_interface:
        db_manager.client_name = client_name

        if client_name not in st.session_state.vector_stores:
            vector_store = db_manager.get_or_create_collection()
            st.session_state.vector_stores[client_name] = vector_store
        else:
            vector_store = st.session_state.vector_stores[client_name]

        rag_system = RAGSystem(vector_store)
        st.sidebar.success(f"Active client: {client_name}")

        # ------------------------
        # Upload Documents
        # ------------------------
        st.sidebar.subheader("Upload Documents")
        uploaded_files = st.sidebar.file_uploader(
            "Upload documents", type=["pdf", "txt", "docx"], accept_multiple_files=True
        )

        if st.sidebar.button("Add Documents"):
            if uploaded_files:
                client_data_dir = os.path.join("data", client_name)
                os.makedirs(client_data_dir, exist_ok=True)

                paths = []
                for f in uploaded_files:
                    save_path = os.path.join(client_data_dir, f.name)
                    with open(save_path, "wb") as out:
                        out.write(f.getbuffer())
                    paths.append(save_path)

                db_manager.add_documents(paths)
                st.sidebar.success(f"Added {len(paths)} documents for client '{client_name}'")
            else:
                st.sidebar.warning("Please select at least one document.")

        # ------------------------
        # List Documents per client
        # ------------------------
        with st.sidebar.expander(f"📁 {client_name} documents", expanded=False):
            documents = db_manager.list_documents(client_name)

            if documents:
                for doc in documents:
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.markdown(f"**📄 {os.path.basename(doc['source'])}**")
                        st.caption(doc["content"][:120].replace("\n", " ") + "...")
                    with col2:
                        if st.button("🗑️", key=f"delete_{doc['source']}"):
                            result = db_manager.delete_document(client_name, doc["source"])
                            st.success(result)
                            st.rerun()
            else:
                st.info("No documents found for this client.")





        # ------------------------
        # Query Section
        # ------------------------
        st.header(f"Query Client: {client_name}")
        # 🔹 Safe clearing before widget instantiation
        if st.session_state.clear_chat_trigger:
            st.session_state.query_input = ""
            st.session_state.clear_chat_trigger = False
        query = st.text_input("Enter your question:", key="query_input")
        top_k = st.slider("Number of top documents to retrieve", 1, 10, 3)

        if st.button("Search") and query.strip():
            with st.spinner("Retrieving and generating answer..."):
                result = rag_system.query(query.strip(), top_k)

                st.markdown("### 💡 GPT Answer")
                st.markdown(
                    f"<div style='background-color:#f0f2f6;padding:15px;border-radius:10px'>{result['answer']}</div>",
                    unsafe_allow_html=True
                )

                st.markdown("### 📚 Source Documents")
                if result.get("sources"):
                    for i, src in enumerate(result["sources"], 1):
                        score = src["score"]
                        with st.expander(f"Document {i} - Score: {score}"):
                            st.write(src["document"].page_content)
                else:
                    st.info("No source documents retrieved.")

        # ------------------------
        # Clear Chat
        # ------------------------
        if st.session_state.query_input.strip():
            if st.button("Clear Chat"):
                st.session_state.clear_chat_trigger = True
                st.rerun()
    else:
        st.info("Please select an existing client or add a new client to start.")


if __name__ == "__main__":
    main()



